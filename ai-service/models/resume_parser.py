# models/resume_parser.py
# Phase 1 – Data Ingestion & Parsing
# Supports PDF (via pdfplumber + PyMuPDF fallback) and DOCX

from __future__ import annotations

import io
import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────────────────────
# PDF Extraction
# ──────────────────────────────────────────────────────────────────────────────

def _extract_pdf_pdfplumber(file_bytes: bytes) -> str:
    """Primary PDF extractor using pdfplumber (better table/layout handling)."""
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_pdf_pymupdf(file_bytes: bytes) -> str:
    """Fallback PDF extractor using PyMuPDF (fitz)."""
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text("text"))
    return "\n".join(text_parts)

def _text_quality_score(text: str) -> float:
    """
    Higher score means cleaner extraction.

    Heuristic: reward normal-length alpha tokens and penalize very long joined tokens
    that commonly appear when spaces are dropped during PDF extraction.
    """
    if not text or not text.strip():
        return float("-inf")

    tokens = re.findall(r"[A-Za-z]+", text)
    if not tokens:
        return float("-inf")

    normal = sum(1 for t in tokens if 2 <= len(t) <= 16)
    long_joined = sum(1 for t in tokens if len(t) >= 20)
    ultra_joined = sum(1 for t in tokens if len(t) >= 30)

    # Penalize extraction artifacts strongly.
    return float(normal - (long_joined * 3) - (ultra_joined * 6))


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract raw text from PDF bytes.
    Runs both extractors when possible and picks the cleaner output.
    """
    plumber_text = ""
    pymupdf_text = ""

    try:
        plumber_text = _extract_pdf_pdfplumber(file_bytes)
    except Exception as e:
        logger.warning("pdfplumber failed: %s", e)

    try:
        pymupdf_text = _extract_pdf_pymupdf(file_bytes)
    except Exception as e:
        logger.warning("PyMuPDF failed: %s", e)

    if not plumber_text.strip() and not pymupdf_text.strip():
        logger.error("Both PDF extractors returned empty text.")
        return ""

    if not plumber_text.strip():
        return pymupdf_text
    if not pymupdf_text.strip():
        return plumber_text

    plumber_score = _text_quality_score(plumber_text)
    pymupdf_score = _text_quality_score(pymupdf_text)

    if pymupdf_score > plumber_score:
        logger.info("Using PyMuPDF extraction (quality %.2f > %.2f).", pymupdf_score, plumber_score)
        return pymupdf_text

    return plumber_text


# ──────────────────────────────────────────────────────────────────────────────
# DOCX Extraction
# ──────────────────────────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from DOCX bytes using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

    # Also grab text from tables (common in resumes)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in paragraphs:
                    paragraphs.append(cell_text)

    return "\n".join(paragraphs)


# ──────────────────────────────────────────────────────────────────────────────
# Unified Entry Point
# ──────────────────────────────────────────────────────────────────────────────

def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Parse a resume file (PDF or DOCX) and return raw extracted text.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used to detect file type).

    Returns:
        Extracted raw text string.

    Raises:
        ValueError: If file type is unsupported.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif ext in {".docx", ".doc"}:
        text = extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            "Please upload a PDF, DOCX, or TXT file."
        )

    if not text.strip():
        logger.warning("Extracted text is empty for file: %s", filename)

    return text
